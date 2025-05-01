import datetime
import glob
import os
import platform
import re
import shutil
import tempfile
import textwrap
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import openai
import tiktoken
import yt_dlp
from langchain.text_splitter import (
    CharacterTextSplitter,
    MarkdownHeaderTextSplitter
)
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langdetect import LangDetectException, detect
from pydub import AudioSegment


# Настройка пути к ffmpeg
def setup_ffmpeg_path():
    """
    Настраивает пути к FFmpeg в зависимости от платформы.
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))
    if platform.system() == "Windows":
        ffmpeg_bin = os.path.join(current_dir, "ffmpeg.exe")
        ffprobe_bin = os.path.join(current_dir, "ffprobe.exe")
        os.environ["FFMPEG_BINARY"] = ffmpeg_bin
        os.environ["FFPROBE_BINARY"] = ffprobe_bin
        # Для pydub
        AudioSegment.converter = ffmpeg_bin
        AudioSegment.ffmpeg = ffmpeg_bin
        AudioSegment.ffprobe = ffprobe_bin


# Вызываем настройку путей
setup_ffmpeg_path()


# Информация об аудио файлe
def audio_info(audio_file: str) -> AudioSegment:
    """
    Получает информацию об аудио файле.
    
    Args:
        audio_file: Путь к аудио файлу
    
    Returns:
        AudioSegment объект
    """
    audio = AudioSegment.from_file(audio_file)
    print(f'\nПродолжительность: {audio.duration_seconds / 60:.2f} мин.')
    print(f'Частота дискретизаци: {audio.frame_rate}')
    print(f'Количество каналов: {audio.channels}')
    return audio


# Транскрибация аудио в текст (OpenAI - whisper)
def transcribe_audio_whisper(
    audio_path: str,
    file_title: str,
    save_folder_path: str,
    max_duration: int = 10*60*1000
) -> Tuple[str, str]:
    """
    Транскрибация аудиофайла по частям с использованием OpenAI Whisper API.

    Args:
        audio_path: Путь к аудио файлу
        file_title: Название файла для сохранения результатов
        save_folder_path: Папка для сохранения результатов
        max_duration: Максимальная длительность фрагмента (в миллисекундах)

    Returns:
        Кортеж из текста транскрипции и языка транскрибации
    """
    # Создание папки для сохранения результатов, если она ещё не существует
    os.makedirs(save_folder_path, exist_ok=True)

    # Загрузка аудиофайла
    audio = AudioSegment.from_file(audio_path)

    # Создание временной папки для хранения аудио фрагментов
    temp_dir = tempfile.mkdtemp()

    # Инициализация переменных для обработки аудио фрагментов
    current_start_time = 0  # Текущее время начала фрагмента
    chunk_index = 1         # Индекс текущего фрагмента
    transcriptions = []     # Список для хранения всех транскрибаций
    detected_language = None

    # Обработка аудиофайла частями
    while current_start_time < len(audio):
        # Выделение фрагмента из аудиофайла
        chunk = audio[current_start_time:current_start_time + max_duration]
        # Формирование имени и пути файла фрагмента
        chunk_name = f"chunk_{chunk_index}.mp3"
        chunk_path = os.path.join(temp_dir, chunk_name)
        # Экспорт фрагмента
        chunk.export(chunk_path, format="mp3")

        # Проверка размера файла фрагмента на соответствие лимиту API
        if os.path.getsize(chunk_path) > 26000000:  # почти 25 MB
            print(
                f"Фрагмент {chunk_index} превышает максимальный размер для API. "
                "Пробуем уменьшить..."
            )
            max_duration = int(max_duration * 0.8)  # Уменьшение длительности
            os.remove(chunk_path)  # Удаление фрагмента, превышающего лимит
            continue

        # Открытие файла фрагмента для чтения в двоичном режиме
        with open(chunk_path, "rb") as src_file:
            print(f"Транскрибация {chunk_name}...")
            try:
                # Запрос на транскрибацию фрагмента с использованием модели Whisper
                transcript_response = openai.audio.transcriptions.create(
                    model="whisper-1",
                    file=src_file
                )
                
                # Добавление результата транскрибации в список транскрипций
                transcriptions.append(transcript_response.text)
                
                # Пытаемся определить язык от API, если он доступен
                response_language = getattr(transcript_response, 'language', None)
                
                # Сохраняем язык транскрибации от первого фрагмента
                if detected_language is None:
                    # Если API не вернул язык, пробуем определить самостоятельно
                    if not response_language or response_language == "unknown":
                        # Пробуем определить язык самостоятельно из текста
                        text_language = detect_language(transcript_response.text)
                        detected_language = text_language
                    else:
                        detected_language = response_language
                    
                    print(f"Определен язык: {detected_language}")
                    
            except openai.BadRequestError as e:
                print(f"Произошла ошибка: {e}")
                break
                
        # Переход к следующему фрагменту
        current_start_time += max_duration
        chunk_index += 1

    # Сохранение всех транскрибаций в один текстовый файл
    result_text = "\n".join(transcriptions)
    result_path = os.path.join(save_folder_path, f"{file_title}.txt")
    with open(result_path, "w", encoding="utf-8") as f:
        f.write(result_text)
    print(f"Транскрипция сохранена в {result_path}")

    # Если язык не был определен, делаем финальную попытку
    if not detected_language or detected_language == "unknown":
        detected_language = detect_language(result_text)
        print(f"Язык определен из полного текста: {detected_language}")

    # Удаляем временную папку и все файлы в ней
    shutil.rmtree(temp_dir)
    return result_text, detected_language


# Функция для форматирования текста по абзацам
def format_text(text: str, width: int = 120) -> str:
    """
    Форматирует текст по абзацам с заданной шириной.
    
    Args:
        text: Исходный текст
        width: Максимальная ширина строки
    
    Returns:
        Отформатированный текст
    """
    # Разделяем текст на абзацы
    paragraphs = text.split('\n')
    # Форматируем каждый абзац отдельно
    formatted_paragraphs = []
    for paragraph in paragraphs:
        # Используем textwrap.fill для форматирования абзаца
        formatted_paragraph = textwrap.fill(paragraph, width)
        formatted_paragraphs.append(formatted_paragraph)
    # Объединяем абзацы с символом новой строки
    return '\n'.join(formatted_paragraphs)


# Функция возвращает количество токенов в строке в зависимости от модели
def num_tokens_from_string(string: str, model: str = 'gpt-4o-mini') -> int:
    """
    Подсчитывает количество токенов в строке для указанной модели.
    
    Args:
        string: Исходный текст
        model: Модель для подсчета токенов
    
    Returns:
        Количество токенов
    """
    # Получаем имя кодировки для указанной модели
    try:
        encoding_name = tiktoken.encoding_for_model(model).name
        # Получаем объект кодировки на основе имени кодировки
        encoding = tiktoken.get_encoding(encoding_name)
        # Кодируем строку и вычисляем количество токенов
        num_tokens = len(encoding.encode(string)) + 10
    except KeyError:
        # Если модель не найдена, используем cl100k_base
        encoding = tiktoken.get_encoding("cl100k_base")
        num_tokens = len(encoding.encode(string)) + 10
    # Возвращаем количество токенов
    return num_tokens


# (CharacterTextSplitter) Формируем чанки из текста по количеству символов
def split_text(
    text: str,
    chunk_size: int = 30000,
    chunk_overlap: int = 1000
) -> List[str]:
    """
    Разделяет текст на фрагменты (чанки) заданного размера.
    
    Args:
        text: Исходный текст
        chunk_size: Максимальный размер чанка в символах
        chunk_overlap: Количество символов перекрытия между чанками
    
    Returns:
        Список текстовых фрагментов
    """
    # Удалить пустые строки и лишние пробелы
    text = re.sub(r'\s+', ' ', text).strip()
    # Создаем экземпляр CharacterTextSplitter с заданными парамаетрами
    splitter = CharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separator=" "
    )
    return splitter.split_text(text)  # список текстовых чанков


# (MarkdownHeaderTextSplitter) Формируем чанки в формат LangChain Document
def split_markdown_text(markdown_text: str, strip_headers: bool = False):
    """
    Разделяет markdown текст на фрагменты по заголовкам.
    
    Args:
        markdown_text: Исходный markdown текст
        strip_headers: Удалять ли заголовки из текста чанков
    
    Returns:
        Список документов LangChain Document
    """
    # Определяем заголовки, по которым будем разбивать текст
    headers_to_split_on = [
        ("#", "Header 1"),   # Заголовок первого уровня
        ("##", "Header 2"),  # Заголовок второго уровня
        ("###", "Header 3")  # Заголовок третьего уровня
    ]
    # Создаем экземпляр MarkdownHeaderTextSplitter с заданными заголовками
    markdown_splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on,
        strip_headers=strip_headers
    )
    # Разбиваем текст на чанки в формат LangChain Document
    return markdown_splitter.split_text(markdown_text)


# Функция получения ответа от модели
def generate_answer(
    system: str,
    user: str,
    text: str,
    temp: float = 0.3,
    model: str = 'gpt-4o-mini'
) -> str:
    """
    Получает ответ от модели OpenAI.
    
    Args:
        system: Системное сообщение
        user: Пользовательское сообщение
        text: Текст для анализа
        temp: Температура генерации
        model: Модель для использования
    
    Returns:
        Ответ от модели
    """
    messages = [
        {'role': 'system', 'content': system},
        {'role': 'user', 'content': user + '\n' + text}
    ]
    completion = openai.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temp
    )
    return completion.choices[0].message.content


# Обработка текстовых чанков по очереди
def process_text_chunks(
    text_chunks: List[str],
    system: str,
    user: str
) -> str:
    """
    Обрабатывает список текстовых чанков с помощью модели.
    
    Args:
        text_chunks: Список текстовых чанков
        system: Системное сообщение
        user: Пользовательское сообщение
    
    Returns:
        Обработанный текст
    """
    processed_text = ''
    for chunk in text_chunks:
        # Получение ответа от модели для каждого чанка
        answer = generate_answer(system, user, chunk)
        processed_text += f'{answer}\n\n'  # Добавляем ответ в результат
    return processed_text


# Обработка каждого чанка (документа) для формирования методички
def process_documents(
    save_folder_path: str,
    documents,
    system: str,
    user: str,
    original_filename: str = "transcript",
    target_language: str = "русский"
) -> str:
    """
    Обрабатывает список документов и формирует методичку.
    
    Args:
        save_folder_path: Путь для сохранения результатов
        documents: Список документов
        system: Системное сообщение
        user: Пользовательское сообщение
        original_filename: Имя оригинального файла 
        target_language: Целевой язык для конспекта
    
    Returns:
        Текст методички
    """
    processed_text_for_handbook = ""  # Строка для конкатенации обработанного текста
    
    # Получаем стандартную языковую инструкцию
    language_instruction = get_language_instruction(target_language)
    
    # Для каждого документа обрабатываем отдельно с явным указанием языка
    for document in documents:
        # Усиливаем систему инструкцией языка для каждого отдельного документа
        enhanced_system = (
            f"{system}\n\nЭТО КРАЙНЕ ВАЖНО: {language_instruction}\n"
            f"Весь текст, ВКЛЮЧАЯ ЗАГОЛОВКИ, должен быть только на "
            f"{target_language} языке!"
        )
        
        # Усиливаем запрос инструкцией языка
        enhanced_user = (
            f"{user}\n\nВАЖНО: Весь текст должен быть ТОЛЬКО на "
            f"{target_language} языке! Заголовки и всё содержание должны быть "
            f"на {target_language}!"
        )
        
        # Получаем ответ от модели для каждого документа
        answer = generate_answer(
            enhanced_system,
            enhanced_user,
            document.page_content
        )
        # Добавляем обработанный текст в общую строку
        processed_text_for_handbook += f"{answer}\n\n"
    
    # Записываем полученный текст во временный файл с уникальным именем
    result_path = os.path.join(
        save_folder_path, f'{original_filename}_summary_draft.txt'
    )
    with open(result_path, 'w', encoding='utf-8') as f:
        f.write(processed_text_for_handbook)
    
    return processed_text_for_handbook


# Вспомогательная функция для получения языковой инструкции
def get_language_instruction(target_language: str) -> str:
    """
    Возвращает строгие языковые инструкции для указанного языка
    
    Args:
        target_language: Целевой язык (русский, казахский, английский)
        
    Returns:
        Строка с инструкцией на соответствующем языке
    """
    if target_language.lower() == "казахский":
        return """БАРЛЫҚ МӘТІНДІ ТЕК ҚАНА ҚАЗАҚ ТІЛІНДЕ ЖАЗУ КЕРЕК!
        Басқа тілдерді МҮЛДЕМ қолданбаңыз. 
        Тақырыптар, мәтін мазмұны, бөлімдер - БӘРІ қазақ тілінде болуы МІНДЕТТІ.
        Орыс немесе ағылшын сөздерін араластыруға ТЫЙЫМ САЛЫНҒАН.
        БҰЛ НҰСҚАУЛЫҚТЫ ҚАТАҢ ТҮРДЕ САҚТАУ ҚАЖЕТ!"""
    elif target_language.lower() == "английский":
        return """ALL TEXT MUST BE WRITTEN ONLY IN ENGLISH!
        DO NOT use other languages AT ALL.
        Headings, content, sections - EVERYTHING must be in English ONLY.
        DO NOT mix in Russian or Kazakh words under ANY circumstances.
        THIS INSTRUCTION MUST BE FOLLOWED STRICTLY!"""
    else:  # русский по умолчанию
        return """ВЕСЬ ТЕКСТ ДОЛЖЕН БЫТЬ НАПИСАН ТОЛЬКО НА РУССКОМ ЯЗЫКЕ!
        НЕ используйте другие языки ВООБЩЕ.
        Заголовки, содержание, разделы - ВСЁ должно быть ТОЛЬКО на русском языке.
        НЕ смешивайте с казахскими или английскими словами НИ ПРИ КАКИХ ОБСТОЯТЕЛЬСТВАХ.
        ЭТО УКАЗАНИЕ ДОЛЖНО БЫТЬ СТРОГО СОБЛЮДЕНО!"""


# Создание индексной (векторной) базы из чанков и сохранение на диск
def create_db_index_from_documents_save(chunks_documents, index_name: str, path: str):
    """
    Создает векторную базу из документов и сохраняет на диск.
    
    Args:
        chunks_documents: Список документов
        index_name: Имя для базы
        path: Путь для сохранения
    
    Returns:
        Векторная база FAISS
    """
    # Создаем индексную базу с использованием FAISS
    db_index = FAISS.from_documents(
        chunks_documents,
        OpenAIEmbeddings()
    )
    # Сохраняем индексную базу
    db_index.save_local(
        folder_path=path,  # путь к папке
        index_name=index_name  # имя для индексной базы
    )
    return db_index


# Загрузка векторной базы с диска
def load_db_vector(folder_path_db_index: str, index_name: str):
    """
    Загружает векторную базу с диска.
    
    Args:
        folder_path_db_index: Путь к базе
        index_name: Имя базы
    
    Returns:
        Векторная база FAISS
    """
    return FAISS.load_local(
        allow_dangerous_deserialization=True,  # Разрешает десериализацию
        embeddings=OpenAIEmbeddings(),  # Указывает векторные представления
        folder_path=folder_path_db_index,  # путь к сохраненной векторной базе
        index_name=index_name  # имя сохраненной векторной базы
    )


# Функция запроса и ответа от OpenAI с поиском по векторной базе данных
def generate_db_answer(
    query: str,
    db_index,
    k: int = 3,
    verbose: bool = True,
    model: str = 'gpt-4o-mini',
    temp: float = 0.3
) -> str:
    """
    Генерирует ответ на основе поиска по векторной базе.
    
    Args:
        query: Запрос пользователя
        db_index: Векторная база
        k: Количество используемых чанков
        verbose: Выводить ли найденные чанки
        model: Модель для генерации ответа
        temp: Температура генерации
    
    Returns:
        Ответ модели
    """
    # Поиск чанков по векторной базе данных
    similar_documents = db_index.similarity_search(query, k=k)
    
    # Формирование текстового контента из выбранных чанков для модели
    message_content = re.sub(
        r'\n{2}', ' ', 
        '\n '.join([
            f'Отрывок документа № {i+1}:\n' + doc.page_content
            for i, doc in enumerate(similar_documents)
        ])
    )
    
    if verbose:
        print(message_content)  # печать на экран выбранных чанков
    
    messages = [
        {
            "role": "system",
            "content": f'Ответь подробно на вопрос пользователя на основании '
                      f'информации из базы знаний: \n{message_content}'
        },
        {"role": "user", "content": f'Вопрос пользователя: {query}'}
    ]
    
    response = openai.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temp
    )
    
    return response.choices[0].message.content


# Сохранение текста в формате DOCX
def save_text_to_docx(text: str, file_path: str) -> None:
    """
    Сохраняет текст в файл DOCX.
    
    Args:
        text: Текст для сохранения
        file_path: Путь для сохранения файла
    """
    from docx import Document
    
    # Создаем новый документ
    doc = Document()
    
    # Добавляем текст в документ
    paragraphs = text.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            p = doc.add_paragraph(paragraph)
    
    # Сохраняем документ
    doc.save(file_path)
    print(f"Документ сохранен: {file_path}")


# Сохранение markdown текста в формате DOCX с форматированием
def markdown_to_docx(markdown_text: str, file_path: str) -> None:
    """
    Сохраняет markdown текст в DOCX с форматированием, приближенным к Streamlit.
    
    Args:
        markdown_text: Текст в формате Markdown
        file_path: Путь для сохранения файла
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import re

    # Создаем новый документ
    doc = Document()

    # Устанавливаем поля страницы
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Очищаем markdown от лишних пустых строк и пробелов
    markdown_text = re.sub(r'\s+$', '', markdown_text, flags=re.MULTILINE)
    markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
    lines = [line.rstrip() for line in markdown_text.strip().split('\n')]

    # Удаляем лишние пустые строки между блоками
    cleaned_lines = []
    prev_empty = True
    for line in lines:
        if not line.strip():
            if not prev_empty:
                cleaned_lines.append('')
            prev_empty = True
        else:
            cleaned_lines.append(line)
            prev_empty = False

    # Функция для применения форматирования текста внутри параграфа
    def process_formatted_text(paragraph, text):
        text = text.replace('\\*', '___ASTERISK___')
        formats = [
            (
                r'\*\*\*(.*?)\*\*\*',
                lambda t: {'text': t, 'bold': True, 'italic': True}
            ),
            (r'\*\*(.*?)\*\*', lambda t: {'text': t, 'bold': True}),
            (r'\*(.*?)\*', lambda t: {'text': t, 'italic': True}),
        ]
        tokens = [(text, {})]
        for pattern, formatter in formats:
            new_tokens = []
            for token_text, token_format in tokens:
                if token_format:
                    new_tokens.append((token_text, token_format))
                    continue
                parts = []
                last_end = 0
                for match in re.finditer(pattern, token_text):
                    if match.start() > last_end:
                        parts.append((token_text[last_end:match.start()], {}))
                    format_props = formatter(match.group(1))
                    parts.append((format_props['text'], format_props))
                    last_end = match.end()
                if last_end < len(token_text):
                    parts.append((token_text[last_end:], {}))
                if parts:
                    new_tokens.extend(parts)
                else:
                    new_tokens.append((token_text, token_format))
            tokens = new_tokens
        for token_text, token_format in tokens:
            run = paragraph.add_run(token_text.replace('___ASTERISK___', '*'))
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            if token_format.get('bold'):
                run.bold = True
            if token_format.get('italic'):
                run.italic = True

    # Основной цикл по строкам
    i = 0
    while i < len(cleaned_lines):
        line = cleaned_lines[i]
        # Заголовки
        header_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if header_match:
            level = len(header_match.group(1))
            header_text = header_match.group(2)
            p = doc.add_heading(header_text, level=level)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
            i += 1
            continue
        # Маркированные списки
        bullet_match = re.match(r'^(\s*)[*\-+]\s+(.+)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            list_text = bullet_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2))
            process_formatted_text(p, list_text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            i += 1
            continue
        # Нумерованные списки
        number_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if number_match:
            indent = len(number_match.group(1))
            list_text = number_match.group(2)
            p = doc.add_paragraph(style='List Number')
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2))
            process_formatted_text(p, list_text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            i += 1
            continue
        # Вложенные списки (o, °)
        nested_match = re.match(r'^(\s+)[o°]\s+(.+)', line)
        if nested_match:
            indent = len(nested_match.group(1))
            list_text = nested_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent // 2))
            process_formatted_text(p, list_text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            i += 1
            continue
        # Пустая строка — пропускаем (не добавляем параграф)
        if not line.strip():
            i += 1
            continue
        # Обычный абзац
        p = doc.add_paragraph()
        process_formatted_text(p, line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        i += 1

    # Удаляем пустые параграфы в конце
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    doc.save(file_path)
    print(f"Документ с форматированием Streamlit сохранен: {file_path}")


def format_transcription_paragraphs(text: str, model: str = 'gpt-4o-mini') -> str:
    """
    Форматирует транскрибацию на абзацы с помощью ChatGPT, не изменяя сам текст.
    
    Args:
        text: Исходный текст транскрибации
        model: Модель OpenAI для форматирования
    
    Returns:
        Текст с разбивкой на абзацы
    """
    import openai
    system = (
        "Ты профессиональный редактор. Тебе дан текст транскрибации, "
        "в котором нет абзацев. Разбей его на абзацы так, чтобы текст "
        "выглядел читабельно и удобно для восприятия. Не изменяй и не "
        "сокращай сам текст, только расставь абзацы. Не добавляй ничего от себя."
    )
    user = (
        "Разбей этот текст на абзацы, чтобы он выглядел как связный, "
        "аккуратно оформленный текст. Не меняй и не сокращай сам текст, "
        "только оформи абзацы. Текст:\n" + text
    )
    response = openai.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user}
        ],
        temperature=0.1
    )
    return response.choices[0].message.content.strip()


def translate_text_gpt(text: str, target_language: str, model: str = 'gpt-4o-mini') -> str:
    """
    Переводит текст на целевой язык с помощью GPT-4o-mini.
    
    Args:
        text: Исходный текст
        target_language: Язык перевода ("русский", "казахский", "английский")
        model: Модель OpenAI для перевода
    
    Returns:
        Переведённый текст
    """
    language_map = {
        "русский": "Russian",
        "казахский": "Kazakh",
        "английский": "English"
    }
    lang = language_map.get(target_language.lower(), target_language)
    system = (
        f"Ты профессиональный переводчик. Переведи текст на {lang}. "
        f"Сохрани структуру и смысл. Не добавляй ничего от себя."
    )
    user = f"Переведи на {lang}:\n{text}"
    response = openai.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user}
        ],
        temperature=0.1
    )
    return response.choices[0].message.content.strip()


def detect_language(text: str) -> str:
    """
    Определяет язык текста с большей точностью, используя несколько методов.
    Поддерживает все языки, включая корейский, японский, китайский и другие.
    
    Args:
        text: Текст для анализа
        
    Returns:
        Код языка ("ru", "en", "kk", "ko" и др.)
    """
    # Проверка на пустой текст
    if not text or text.strip() == "":
        return "unknown"
    
    # Использовать первые 1000 символов для более точного определения
    sample_text = text[:1000] if len(text) > 1000 else text
    
    try:
        # Пробуем определить язык с помощью langdetect
        lang_code = detect(sample_text)
        
        # Проверяем язык на некоторые известные проблемы определения
        if lang_code == "ru" and any(eng_word in sample_text.lower() for eng_word in [
            "the", "and", "you", "is", "are", "this", "that", "what", "where",
            "when", "how", "fibonacci", "trend", "level", "market", "usd",
            "uptrend", "continue", "profit"
        ]):
            # Если обнаружены очевидные английские слова, но язык определился как русский
            return "en"
            
        # Определение корейского языка (проверка на наличие корейских символов)
        if any('\uAC00' <= c <= '\uD7A3' for c in sample_text):
            return "ko"
        
        # Определение японского языка (проверка на наличие японских символов)
        if any('\u3040' <= c <= '\u30FF' for c in sample_text):
            return "ja"
            
        # Другие корректировки при необходимости
        if lang_code == "mk" and any(rus_word in sample_text.lower() for rus_word in [
            "это", "привет", "спасибо", "пожалуйста", "да", "нет",
            "говорить", "русский"
        ]):
            # Македонский иногда путается с русским
            return "ru"
            
        return lang_code
    except LangDetectException:
        # В случае ошибки проверяем наличие символов определенных языков
        if any('\uAC00' <= c <= '\uD7A3' for c in sample_text):  # Корейский
            return "ko"
        elif any('\u3040' <= c <= '\u30FF' for c in sample_text):  # Японский
            return "ja"
        elif any('\u4E00' <= c <= '\u9FFF' for c in sample_text):  # Китайский
            return "zh"
        
        # Если не удалось определить, возвращаем unknown
        return "unknown"


def analyze_temp_files(directory):
    """
    Анализирует временные файлы в указанной директории
    
    Args:
        directory: Путь к директории для анализа
        
    Returns:
        tuple: (количество файлов, общий размер в байтах, словарь с файлами по возрасту)
    """
    if not os.path.exists(directory):
        return 0, 0, {}
    
    now = datetime.datetime.now()
    count = 0
    total_size = 0
    files_by_age = {
        "less_than_day": {"count": 0, "size": 0},
        "1_to_7_days": {"count": 0, "size": 0},
        "older_than_7_days": {"count": 0, "size": 0}
    }
    
    for file_path in glob.glob(f"{directory}/**/*", recursive=True):
        if os.path.isfile(file_path):
            try:
                file_size = os.path.getsize(file_path)
                file_mtime = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
                file_age = now - file_mtime
                
                count += 1
                total_size += file_size
                
                # Распределяем по возрастным группам
                if file_age.days < 1:
                    files_by_age["less_than_day"]["count"] += 1
                    files_by_age["less_than_day"]["size"] += file_size
                elif file_age.days < 7:
                    files_by_age["1_to_7_days"]["count"] += 1
                    files_by_age["1_to_7_days"]["size"] += file_size
                else:
                    files_by_age["older_than_7_days"]["count"] += 1
                    files_by_age["older_than_7_days"]["size"] += file_size
            except Exception as e:
                print(f"Ошибка при анализе файла {file_path}: {str(e)}")
    
    return count, total_size, files_by_age


def format_file_size(size_bytes):
    """
    Форматирует размер в байтах в удобочитаемую строку
    
    Args:
        size_bytes: Размер в байтах
        
    Returns:
        str: Отформатированный размер
    """
    # Определяем единицы измерения
    for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
        if size_bytes < 1024 or unit == 'TB':
            return f"{size_bytes:.2f} {unit}"
        size_bytes /= 1024.0